"""将 Danny 面试测试 MD 转为 Word 文档"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 样式 ──
style_normal = doc.styles['Normal']
style_normal.font.name = '微软雅黑'
style_normal.font.size = Pt(10.5)
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.3

for level in range(1, 4):
    style = doc.styles[f'Heading {level}']
    style.font.name = '微软雅黑'
    style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    if level == 1:
        style.font.size = Pt(18)
    elif level == 2:
        style.font.size = Pt(14)
    else:
        style.font.size = Pt(12)

def add_table(doc, header_row, data_rows):
    """添加格式化表格"""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # data
    for r_idx, row_data in enumerate(data_rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacer

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)

def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)

# ═══════════════════════════════════════
# 正文内容
# ═══════════════════════════════════════

doc.add_heading('AI美术面试测试题 — NPC角色立绘制作', level=1)

doc.add_heading('测试说明', level=2)
p = doc.add_paragraph()
p.add_run('本测试用于评估候选人使用AI美术工具（Midjourney / Stable Diffusion / ComfyUI 等）进行游戏角色立绘制作的能力。请在 ').font.size = Pt(10.5)
run = p.add_run('3小时')
run.bold = True
run.font.size = Pt(10.5)
p.add_run(' 内完成以下全部交付物。').font.size = Pt(10.5)

p = doc.add_paragraph()
run = p.add_run('评估维度：')
run.bold = True
add_bullet(doc, '角色理解力：能否从文字设定中提炼视觉特征')
add_bullet(doc, '风格还原力：能否匹配已有角色的美术风格')
add_bullet(doc, '工具熟练度：AI出图效率、prompt工程、后期修图能力')
add_bullet(doc, '一致性控制：同一角色多张图的面部/服装/气质一致性')
add_bullet(doc, '审美判断力：构图、光影、色彩的综合审美水平')

# ── 项目背景 ──
doc.add_heading('项目背景', level=2)
add_bullet(doc, '2D侦探推理AVG（视觉小说）', bold_prefix='游戏类型：')
add_bullet(doc, '1925年，美国芝加哥，禁酒令时期', bold_prefix='时代背景：')
add_bullet(doc, '半写实绘画风，偏欧美graphic novel质感，色调偏暗沉/去饱和，带有noir（黑色电影）氛围', bold_prefix='美术风格：')
add_bullet(doc, 'Unity 2D，角色以立绘形式呈现于对话界面', bold_prefix='引擎：')

doc.add_heading('现有角色风格参考', level=3)
p = doc.add_paragraph('以下为已完成的同项目角色立绘，')
run = p.add_run('你的产出必须在风格上与这些角色保持统一')
run.bold = True
p.add_run('：')

add_table(doc,
    ['参考角色', '文件', '特征说明'],
    [
        ['Morrison（腐败警探）', '参考图/morrison_big.png', '中年男性，棕色警服外套，严肃凶狠，室内暖光'],
        ['Tommy（俱乐部经理）', '参考图/tommy_big.png', '瘦削男性，白衬衫+酒红领结，精明世故，奢华内景'],
        ['Jimmy（酒保）', '参考图/jimmy_big.png', '年轻男性，白衬衫+围裙，疲惫忧郁，朴素背景'],
    ])

doc.add_heading('风格关键词提取（供参考）', level=3)
add_bullet(doc, '油画质感，可见笔触肌理，非照片级写实', bold_prefix='笔触：')
add_bullet(doc, '强对比，单侧主光源，noir风格明暗', bold_prefix='光影：')
add_bullet(doc, '低饱和度，以棕/灰/橄榄绿为主色系', bold_prefix='色调：')
add_bullet(doc, '胸部以上半身像，3/4侧面角度为主', bold_prefix='构图：')
add_bullet(doc, '简化的场景暗示（模糊处理），不喧宾夺主', bold_prefix='背景：')
add_bullet(doc, '五官写实但略带风格化，线条感强', bold_prefix='面部：')

# ── 角色设定 ──
doc.add_heading('角色设定：Danny Kowalski', level=2)

doc.add_heading('基本信息', level=3)
add_table(doc,
    ['项目', '内容'],
    [
        ['姓名', 'Danny Kowalski'],
        ['年龄', '35岁（1890年生）'],
        ['种族', '波兰裔白人'],
        ['职业', '工厂工人（消极怠工，混日子）'],
        ['居住', '寄住在叔叔Frank家中（芝加哥南区工人社区）'],
        ['身份', 'Frank的侄子，从乡下投奔叔叔，是Frank唯一的亲人'],
        ['叙事角色', '次要反派——因遗产纠纷诬陷他人，被玩家指证揭穿'],
    ])

doc.add_heading('外貌特征（必须体现）', level=3)
add_table(doc,
    ['维度', '描述', '设计意图'],
    [
        ['体型', '中等身材，略显臃肿，肚子微凸', '好吃懒做、缺乏锻炼的生活方式'],
        ['面容', '35岁但看起来更老，面色不佳，眼袋明显', '不良嗜好+懒散生活的痕迹'],
        ['眼神', '飘忽不定，带有急躁和贪婪感', '核心性格外化——寄生虫式的不安与觊觎'],
        ['发型', '棕色短发，油腻凌乱，未经打理', '邋遢、不修边幅'],
        ['面部特征', '下颌略宽，法令纹深，嘴角习惯性下撇', '常年不满和怨气的积累'],
        ['肤色', '偏白但不健康，略显蜡黄', '工厂工人但不爱运动的气色'],
    ])

doc.add_heading('服装设计（必须体现）', level=3)
add_table(doc,
    ['部件', '描述', '设计意图'],
    [
        ['上衣', '灰蓝色或土黄色工人衬衫，袖口卷起但歪斜不齐', '工厂工人身份，但邋遢随意'],
        ['外搭', '深色旧背心或旧马甲，扣子缺一个', '不体面但勉强"穿了件外套"的感觉'],
        ['裤子', '深色工装裤，膝盖处有磨损（如可见）', '底层劳动者'],
        ['整体感', '衣服皱巴巴，领口松垮，像是随手套上的', '与Tommy的精致、Morrison的整肃形成对比'],
    ])

doc.add_heading('气质关键词', level=3)
p = doc.add_paragraph()
run = p.add_run('必须传达的核心气质（按优先级）：')
run.bold = True
add_numbered(doc, ' — 整个人散发出"得过且过"的气息，姿态松垮', bold_prefix='懒散寄生')
add_numbered(doc, ' — 眼神中有藏不住的觊觎和不满', bold_prefix='贪婪不甘')
add_numbered(doc, ' — 不是恶人的精明，是小人物的粗糙和短视', bold_prefix='底层粗鄙')
add_numbered(doc, ' — 随时可能爆发的焦躁感', bold_prefix='急躁冲动')

p = doc.add_paragraph()
run = p.add_run('不应有的气质：')
run.bold = True
add_bullet(doc, '不要帅气/英俊（他不是魅力型反派）')
add_bullet(doc, '不要精明/城府深（他不聪明，只是贪婪）')
add_bullet(doc, '不要可怜兮兮（他可悲但不值得同情的外表）')
add_bullet(doc, '不要过于肮脏/流浪汉化（他有工作，只是邋遢）')

# ── 交付物要求 ──
doc.add_heading('交付物要求', level=2)

doc.add_heading('交付物 A：角色立绘正图（必做）', level=3)
p = doc.add_paragraph()
run = p.add_run('数量：')
run.bold = True
p.add_run('1张定稿')

p = doc.add_paragraph()
run = p.add_run('规格：')
run.bold = True
add_bullet(doc, '尺寸：1024×1024px 或以上')
add_bullet(doc, '构图：胸部以上半身像，3/4侧面（与参考角色一致）')
add_bullet(doc, '背景：简化的工人社区室内暗示（砖墙/旧家具/昏暗灯光）')
add_bullet(doc, '光影：单侧暖光（油灯/白炽灯感），noir风格强对比')

p = doc.add_paragraph()
run = p.add_run('质量标准：')
run.bold = True
add_bullet(doc, '与三张参考角色放在一起，风格不违和')
add_bullet(doc, '面部细节清晰，五官可辨认')
add_bullet(doc, '服装材质可信（棉麻布料质感）')
add_bullet(doc, '无明显AI瑕疵（多余手指、融合伪影、文字乱码等）')

doc.add_heading('交付物 B：表情差分（必做）', level=3)
p = doc.add_paragraph('基于交付物A的面部，制作 ')
run = p.add_run('3个表情变体')
run.bold = True
p.add_run('，要求面部以外的部分（服装、姿态、背景）保持一致：')

add_table(doc,
    ['编号', '表情', '情境描述', '关键表现'],
    [
        ['B1', '默认/日常', '平时懒散的状态', '眼神飘忽，嘴角微撇，漫不经心'],
        ['B2', '愤怒', '听到遗嘱内容，爆发', '瞪眼，龇牙，面部肌肉紧绷，青筋'],
        ['B3', '崩溃/不甘', '被指证后情绪崩溃', '眼眶泛红，表情扭曲，嘴角颤抖，不甘心'],
    ])

p = doc.add_paragraph()
run = p.add_run('规格：')
run.bold = True
p.add_run('每张512×512px 或以上，仅面部特写即可')

doc.add_heading('交付物 C：设计过程说明（必做）', level=3)
p = doc.add_paragraph('提交一份简短的文字说明（300-500字），包含：')
add_numbered(doc, '：使用了哪些AI工具和后期软件', bold_prefix='工具链')
add_numbered(doc, '：生成正图时使用的核心prompt（可脱敏/简化）', bold_prefix='关键Prompt')
add_numbered(doc, '：简述从初稿到定稿的调整思路（附2-3张过程图）', bold_prefix='迭代过程')
add_numbered(doc, '：如何保证表情差分与正图的角色一致性', bold_prefix='一致性方案')
add_numbered(doc, '：如何分析参考图并还原风格', bold_prefix='风格匹配方法')

# ── 评分标准 ──
doc.add_heading('评分标准', level=2)

add_table(doc,
    ['维度', '权重', '评分要点'],
    [
        ['风格一致性', '30%', '与参考角色放在一起是否像"同一个游戏的角色"'],
        ['角色表现力', '25%', '是否一眼能感受到"懒散、贪婪、粗鄙"的气质'],
        ['技术质量', '20%', '无AI瑕疵、细节清晰、光影合理、材质可信'],
        ['表情差分', '15%', '三个表情是否到位、面部一致性是否良好'],
        ['过程说明', '10%', '工作流是否清晰专业、是否展现了可复制的方法论'],
    ])

doc.add_heading('加分项（非必需）', level=3)
add_bullet(doc, '提供额外的全身像或动态姿势变体')
add_bullet(doc, '表情差分超过3个（如增加"诬陷时指着别人"的姿态）')
add_bullet(doc, '提供可编辑的PSD/分层文件')
add_bullet(doc, '展示ControlNet / IP-Adapter等高级一致性技术的使用')

doc.add_heading('扣分项', level=3)
add_bullet(doc, '与参考角色风格明显不匹配（如日系二次元风、照片级写实等）')
add_bullet(doc, '角色气质与设定严重偏离（如画成了帅气青年或凶狠恶棍）')
add_bullet(doc, '表情差分中角色面部特征不一致（换了一张脸）')
add_bullet(doc, 'AI伪影未清理（多余手指、文字噪点、面部融合等）')
add_bullet(doc, '未提供过程说明')

# ── 提交方式 ──
doc.add_heading('提交方式', level=2)
p = doc.add_paragraph('请将所有文件打包为ZIP，按以下结构组织：')

code_text = """Danny_Kowalski_[姓名]/
├── A_立绘正图/
│   └── danny_final.png
├── B_表情差分/
│   ├── danny_default.png
│   ├── danny_angry.png
│   └── danny_collapse.png
├── C_过程说明/
│   ├── 设计说明.pdf (或 .md / .docx)
│   └── 过程图/
│       ├── iteration_01.png
│       └── iteration_02.png
└── D_加分项/  (可选)
    └── ..."""

p = doc.add_paragraph()
run = p.add_run(code_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# ── 页脚说明 ──
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('本测试题仅用于面试评估，测试产出物不会用于商业用途。')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
p = doc.add_paragraph()
run = p.add_run('如有疑问请联系面试对接人。')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── 保存 ──
out_path = Path(__file__).parent / '面试测试_Danny角色立绘需求.docx'
doc.save(str(out_path))
print(f'Done: {out_path}')
