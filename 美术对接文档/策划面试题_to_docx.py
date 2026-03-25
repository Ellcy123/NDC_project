# -*- coding: utf-8 -*-
"""策划面试题精简版 -> Word"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

doc = Document()

# ── 页面 ──
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

# ── 样式 ──
sn = doc.styles['Normal']
sn.font.name = '微软雅黑'
sn.font.size = Pt(10.5)
sn.paragraph_format.space_after = Pt(4)
sn.paragraph_format.line_spacing = 1.35

for lv in range(1, 4):
    st = doc.styles[f'Heading {lv}']
    st.font.name = '微软雅黑'
    st.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    st.font.size = [None, Pt(18), Pt(14), Pt(12)][lv]

GRAY = RGBColor(0x66, 0x66, 0x66)
BLUE = RGBColor(0x1A, 0x5C, 0x8A)

def add_q(doc, label, text):
    """添加问题"""
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLUE
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)

def add_followup(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GRAY

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.5)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = val
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    doc.add_paragraph()

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('─' * 60)
    r.font.size = Pt(6)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ═══════════════════════════════════════
# 正文
# ═══════════════════════════════════════

doc.add_heading('叙事推理策划 - 面试问题单', level=1)
add_note(doc, '岗位：叙事推理方向游戏策划  |  建议时长：60-90min')

add_divider(doc)

# ── 一 ──
doc.add_heading('一、经验摸底（15min）', level=2)

add_q(doc, 'Q1. 项目经历',
      '"简单介绍一下自己，重点聊你参与过的最完整的一个项目，你在里面具体负责什么。"')

add_q(doc, 'Q2. 推理游戏分析',
      '"你玩过的推理游戏里，哪一部你觉得证据设计做得最好？好在哪里？"')
add_followup(doc, '追问：能具体到哪一关、哪条证据链吗？')

add_divider(doc)

# ── 二 ──
doc.add_heading('二、专业能力（25min）', level=2)

add_q(doc, 'Q3. 证据链设计',
      '"假设你要设计一个推理章节，核心真相是\'A杀了B\'。你会怎么设计证据链，让玩家自己推出这个结论，而不是直接告诉他？"')
add_followup(doc, '追问：如果玩家收集证据的顺序不同，怎么保证推理体验不崩？')
add_followup(doc, '追问：干扰项怎么设计？原则是什么？')
add_followup(doc, '追问：怎么区分"合理的难度"和"不公平的谜题"？')

add_q(doc, 'Q4. 叙事反转',
      '"你觉得一个好的叙事反转需要满足什么条件？举一个你印象深刻的反转，分析它为什么有效。"')

add_q(doc, 'Q5. 一致性维护',
      '"如果让你维护一个有6个章节、20+个NPC、50+条证据的项目，你怎么保证不出现逻辑矛盾？"')
add_followup(doc, '追问：一条证据在第2章和第5章都引用了，第2章改了措辞，怎么确保第5章不矛盾？')

add_q(doc, 'Q6. 难度判断',
      '"你在做策划时，怎么判断一个谜题对玩家来说太简单还是太难？"')

add_divider(doc)

# ── 三 ──
doc.add_heading('三、实操题（20-30min）', level=2)
add_note(doc, '以下两题任选其一，或视时间两题都做。打印给候选人阅读后作答。')

doc.add_heading('实操A：证据链分析', level=3)

p = doc.add_paragraph()
r = p.add_run('案情：')
r.bold = True
p.add_run('酒吧老板Webb在自己办公室被人枪杀。警探Morrison声称巡逻时听到枪声，破门而入发现尸体，立即报告。')

p = doc.add_paragraph()
r = p.add_run('现有证据：')
r.bold = True

for i, ev in enumerate([
    'Webb桌上有一块沾湿的毛巾（分析后发现含有氯仿）',
    '酒保Jimmy证词："Morrison在枪声前10分钟就已经到了酒吧"',
    'Morrison报告写着"22:15听到枪声"',
    'Webb身上有拖拽痕迹',
    'Morrison车后备箱有手套和绳子',
    'Morrison有赌债，欠高利贷3000美元',
], 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(ev).font.size = Pt(10.5)

p = doc.add_paragraph()
r = p.add_run('问题：')
r.bold = True
for q in [
    '这些证据能直接证明"Morrison杀了Webb"吗？为什么？',
    '还缺什么关键证据？列出至少2个需要补充的方向。',
    '如果做成3轮逐步揭示的结构，每轮分别揭示什么？',
]:
    p = doc.add_paragraph(style='List Number')
    p.add_run(q).font.size = Pt(10.5)

doc.add_paragraph()

# 实操B
doc.add_heading('实操B：对话设计找茬', level=3)

p = doc.add_paragraph()
r = p.add_run('背景：')
r.bold = True
p.add_run('第2轮循环，酒保Jimmy接受侦探询问。请找出以下对话的设计问题（至少3个）。')

doc.add_paragraph()

dialogue = [
    ('Zack', '昨晚酒吧打烊是什么时候？'),
    ('Jimmy', '和平时一样，凌晨2点关门。不过Webb老板昨晚走得早，大概10点左右就回办公室了。'),
    ('Zack', 'Morrison警探什么时候来的？'),
    ('Jimmy', 'Morrison？他经常来。昨晚大概10点出头就到了，在吧台喝了杯威士忌，然后就上楼去找Webb了。不过说来也奇怪，他当时带了一个黑色大包。'),
    ('Zack', '所以Morrison在枪声之前就已经到了酒吧。他带着包去找Webb，很可能事先就做好了准备。看来Morrison就是凶手。'),
    ('Jimmy', '我...我不知道，我只是个酒保...Morrison警探平时对我还不错。'),
    (None, '[获得证据：Morrison的赌债记录]'),
    ('Zack', '我还需要了解一下Webb的生意情况。'),
    ('Jimmy', 'Webb? 他做很多生意，不都是合法的。他搞私酒、放高利贷、还有一些古董买卖。其实Morrison就欠了Webb的高利贷3000块，我听Webb亲口说的，他打算用这笔债来要挟Morrison帮他办事。'),
    (None, '[获得证据：Webb的高利贷生意]'),
    (None, '[获得证据：Morrison欠Webb赌债]'),
]

for speaker, line in dialogue:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)
    if speaker is None:
        r = p.add_run(line)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x88, 0x44, 0x00)
    else:
        r = p.add_run(f'{speaker}：')
        r.bold = True
        r.font.size = Pt(10.5)
        r = p.add_run(f'"{line}"')
        r.font.size = Pt(10.5)

doc.add_paragraph()

add_divider(doc)

# ── 四 ──
doc.add_heading('四、AI工具能力（10min）', level=2)

add_q(doc, 'Q7. 使用经历',
      '"你在实际工作中怎么用AI工具的？给我一个最典型的使用场景。"')
add_followup(doc, '追问：AI生成的内容有逻辑错误，你怎么发现和修正？')
add_followup(doc, '追问：策划工作中哪些环节适合用AI，哪些不适合？')

add_q(doc, 'Q8. Prompt能力（选问）',
      '"如果给你一个NPC的设定文档，让你用AI生成对话初稿，你的prompt大概怎么写？说一下结构。"')

add_divider(doc)

# ── 五 ──
doc.add_heading('五、文化适配（5min）', level=2)

add_q(doc, 'Q9. 工作方式',
      '"我们是小团队，策划要自己写文档、查逻辑、跟进度。你习惯什么样的工作方式？"')

add_q(doc, 'Q10. 反向提问',
      '"你有什么想问我们的吗？"')

add_divider(doc)

# ── 评分表 ──
doc.add_heading('评分表', level=2)

add_table(doc,
    ['维度', '权重', '1', '2', '3', '4', '5', '对应题目'],
    [
        ['推理游戏理解深度',   '20%', '', '', '', '', '', 'Q2, Q4'],
        ['证据链/逻辑设计',    '25%', '', '', '', '', '', 'Q3, 实操题'],
        ['叙事节奏与信息控制',  '15%', '', '', '', '', '', 'Q3, Q4, 实操题'],
        ['找漏洞能力',         '15%', '', '', '', '', '', 'Q5, Q6, 实操B'],
        ['AI工具能力',         '10%', '', '', '', '', '', 'Q7, Q8'],
        ['文档与工作习惯',     '10%', '', '', '', '', '', 'Q5, Q9'],
        ['团队适配',           '5%',  '', '', '', '', '', 'Q9, Q10'],
    ])

p = doc.add_paragraph()
r = p.add_run('综合评分：________ / 5')
r.bold = True
r.font.size = Pt(11)

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('录用建议：')
r.bold = True
p.add_run('  \u25a1 推荐录用    \u25a1 可考虑    \u25a1 不推荐')

p = doc.add_paragraph()
r = p.add_run('一票否决：')
r.bold = True
p.add_run('  \u25a1 实操题完全无法完成    \u25a1 无推理游戏经验    \u25a1 AI工具完全不会用')

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('面试官签字：________________    日期：________________').font.size = Pt(10.5)

# ── 保存 ──
out = Path(__file__).parent / '策划面试题_叙事推理方向.docx'
doc.save(str(out))
print(f'Done: {out.name}')
